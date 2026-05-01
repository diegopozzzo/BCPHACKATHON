from __future__ import annotations

import io
import re
from dataclasses import dataclass

from docx import Document
from pypdf import PdfReader


@dataclass(frozen=True, slots=True)
class CVDocument:
    filename: str
    mimetype: str
    text: str


def _clean_text(t: str) -> str:
    t = t.replace("\x00", " ").replace("\u200b", " ")
    t = re.sub(r"[ \t]+", " ", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()


def parse_pdf(data: bytes) -> str:
    r = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in r.pages:
        txt = page.extract_text() or ""
        if txt.strip():
            parts.append(txt)
    return _clean_text("\n\n".join(parts))


def parse_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text)
    return _clean_text("\n".join(parts))


def parse_cv_bytes(*, filename: str, mimetype: str, data: bytes) -> CVDocument:
    fn = (filename or "").lower().strip()
    mt = (mimetype or "").lower().strip()

    if mt == "application/pdf" or fn.endswith(".pdf"):
        text = parse_pdf(data)
        return CVDocument(filename=filename or "cv.pdf", mimetype=mimetype or "application/pdf", text=text)

    if mt in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ) or fn.endswith((".docx", ".doc")):
        text = parse_docx(data)
        return CVDocument(
            filename=filename or "cv.docx",
            mimetype=mimetype
            or "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            text=text,
        )

    # Unknown: best effort as UTF-8 text
    try:
        text = data.decode("utf-8", errors="replace")
    except Exception:  # noqa: BLE001
        text = ""
    return CVDocument(filename=filename or "cv.bin", mimetype=mimetype or "application/octet-stream", text=_clean_text(text))

